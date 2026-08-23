"""
DOCX report building service using python-docx and docxcompose.
Handles creation of item files and final merged reports.

Path strategy
─────────────
The backend runs on Windows.  All file I/O uses the
DATA_* paths defined in .env.
(e.g. C:\\AppIO\\sec10k\\part2),
"""
import os
import re
import tomllib
from pathlib import Path
import tomli_w
from bs4 import BeautifulSoup
from docx import Document
from docxcompose.composer import Composer
from htmldocx import HtmlToDocx

from app.schemas.schemas import BalanceSheetRow, IncomeStatementRow, CashFlowRow, StockholdersEquityRow
from app.core.config import settings
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import logging

logger = logging.getLogger(__name__)

def _ensure_dir(path: str):
    os.makedirs(path, exist_ok=True)


def _q10_base(quarter: int) -> str:
    """
    Quarterly 10-Q base directory, e.g. "sec10q-q3" — see
    env-10q-123.docx. Each quarter now has its own directory tree
    (BASE_10Q + quarter, containing its own intro/part1/part2/sigs
    subdirectories) instead of one shared sec10q directory, so past
    quarters' item files are preserved rather than overwritten by the
    next quarter's run. Every 10-Q path in this file that used to be
    os.path.join(settings.DATA_10Q_PARTx, ...) is now
    os.path.join(_q10_base(quarter), settings.DATA_10Q_PARTx, ...).
    """
    return settings.BASE_10Q + str(quarter)


def _add_heading(doc: Document, text: str, level: int = 2):
    heading_style = doc.styles['Heading 3']
    heading_style.font.bold = True
    heading_style.font.size = Pt(12)
    doc.add_heading(text, level=level)

def _fmt(value: float) -> str:
    """Format exactly as the frontend fmtNum(): negative → (1,234)  zero → —"""
    if value is None or value == 0:
        return "\u2014"
    abs_val = f"{abs(value):,.0f}"
    return f"({abs_val})" if value < 0 else abs_val

def _fmt_abs(value: float) -> str:
    """Absolute-value format — for expense rows stored negative, shown positive."""
    if value is None or value == 0:
        return "\u2014"
    return f"{abs(value):,.0f}"

def _fmt_neg(value: float) -> str:
    """
    Always show in parens (unless zero) regardless of the value's stored
    sign. For expense rows whose current_period is stored POSITIVE (so that
    section-total subtraction math like `revenue_sum - expense_sum` and
    `pretax_sum - tax_sum` keeps working elsewhere in this file), but which
    still need a visual "this is being subtracted" cue — e.g. Interest
    Expense sitting in the same "OTHER INCOME / (EXPENSES)" section as
    Interest Income (shown signed), or Income Tax Expense reducing Net
    Income. Using _fmt_abs on these hid the parens entirely, since the
    stored value was never negative to begin with.
    """
    if value is None or value == 0:
        return "\u2014"
    return f"({abs(value):,.0f})"

def _income_stmt_label(r) -> str:
    """
    Display label for an income-statement row. Tax rows show as
    "Income tax" (SEC/GAAP convention -- a plain deduction line, not a
    parenthesized netted item) regardless of whatever literal acct_name
    ("Income Tax Expense") sits in the chart of accounts, so the CoA
    itself doesn't need to change. xbrl_tagger.py's GAAP_MAPPINGS carries
    "income tax" as an alias of the existing "income tax expense" entry
    so this rename doesn't break tagging's label-text lookup.
    """
    if _classify(r) == "tax":
        return "Income tax"
    return r.acct_name


def _add_dollar(s: str) -> str:
    """
    Prepend a $ to an already-_fmt()-formatted string — e.g. "37,942,448"
    -> "$37,942,448", "(433,455)" -> "$(433,455)". Standard GAAP financial-
    statement convention: only the FIRST data row and the LAST (grand
    total) row of each statement's dollar column carry the $ sign; every
    row in between is unmarked. Leaves the em-dash zero-placeholder alone.
    """
    return s if s == "\u2014" else f"${s}"

# ── Low-level cell / row helpers ─────────────────────────────────────────────

def _set_cell_text(cell, text: str, bold=False, align_right=False,
                   font_name="Times New Roman", font_pt=10):
    para = cell.paragraphs[0]
    para.clear()
    run = para.add_run(str(text))
    run.bold = bold
    run.font.name = font_name
    run.font.size = Pt(font_pt)
    if align_right:
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _right_align_notes_table_columns(doc, label_cols: int = 1):
    """
    Right-align every column after the first `label_cols` column(s), on
    every table HtmlToDocx just parsed out of a notes_agent.py narrative.

    notes_agent.py's _build_*_table_html() functions (segment revenue,
    debt maturity, debt rate breakdown, lease maturity) emit plain
    <td>/<th> markup with no align/style attributes at all - HtmlToDocx
    doesn't reliably map CSS alignment onto python-docx paragraph
    alignment, so every cell defaulted to left-aligned, including the
    numeric columns. That made the numbers hard to visually compare down
    a column.

    Every one of those table-builders has the same shape - column 0 is
    always a text label (fiscal year, category, maturity bucket, interest
    rate band), every column after that is a dollar amount - and
    notes_agent.py instructs the LLM to insert these as an exact
    placeholder token rather than author its own tables, so this holds for
    every table these functions produce. Safe to apply unconditionally to
    every table HtmlToDocx creates in a notes narrative.

    Called right after parser.add_html_to_document(narrative, doc), before
    doc.save() - this sets a REAL w:jc on each cell's paragraph, so it
    round-trips correctly through docxcompose's later merge and gets read
    back by _docx_body_to_html_parts()'s _cell_jc()/_JC_MAP into the
    correct "text-align: right" in the final EDGAR HTML, the same
    mechanism already used for the financial-statement tables.
    """
    for table in doc.tables:
        for row in table.rows:
            for col_idx, cell in enumerate(row.cells):
                if col_idx < label_cols:
                    continue
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _add_cell_border(cell, side: str, style="single", size=8):
    """Add a border to one side of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    el = OxmlElement(f"w:{side}")
    el.set(qn("w:val"), style)
    el.set(qn("w:sz"), str(size))
    el.set(qn("w:space"), "0")
    el.set(qn("w:color"), "000000")
    tcBorders.append(el)
    tcPr.append(tcBorders)


# ── Alternating row shading (financial statements only) ──────────────────────
# Real EDGAR filings zebra-stripe financial-statement tables with
# background-color: #ffffff / #f1f2f2, applied per-<td> rather than once on
# the <tr> — mainly for compatibility with the wider ecosystem of EDGAR HTML
# consumers (viewers, screen readers, print-to-PDF pipelines), not because
# row-level CSS is unsupported (it's standards-compliant: an unstyled <td>
# shows its parent <tr>'s background per the CSS2.1 table background-
# painting order). We apply it the same way _add_cell_border() already
# round-trips real border formatting: set real OOXML shading (w:shd) here
# at docx-generation time, then _docx_body_to_html_parts()'s
# _cell_shading_color() reads it back and — since every cell in the row
# always gets the SAME color — coalesces it into a single
# <tr style="background-color:...">, matching what was actually asked for
# without repeating the style on every <td>.
#
# Scoped by construction, not by table-shape detection: _add_row_shading()
# is only ever called from the four financial-statement row-writers below
# (_data_row/_total_row/_section_header_row and the Stockholders' Equity
# writer's own row loop). Notes tables, the TOC, and the cover page are
# built through entirely different code paths that never call this, so
# they're structurally incapable of picking up shading, regardless of what
# their table markup happens to look like.
_ZEBRA_WHITE = "FFFFFF"
_ZEBRA_GRAY  = "F1F2F2"


def _add_cell_shading(cell, hex_color: str):
    """Set one cell's background shading (OOXML w:shd)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_row_shading(row, hex_color: str):
    """Shade every cell in a table row the same color, so
    _cell_shading_color() below finds a uniform color across the row and
    can emit ONE <tr style="..."> instead of per-<td> styles."""
    for cell in row.cells:
        _add_cell_shading(cell, hex_color)


def _zebra_color(table) -> str:
    """
    Alternating fill for the row that was JUST added via table.add_row()
    — call this AFTER add_row(), not before. Row 0 (the header, created by
    _make_table()/filled by _header_row()) is never shaded — it already
    has its own bold + border treatment, so alternation starts fresh at
    the first data row: 1st data row -> white, 2nd -> gray, 3rd -> white,
    etc. (data_row_idx 0-based: 0, 2, 4... -> white; 1, 3, 5... -> gray).
    """
    data_row_idx = len(table.rows) - 2  # row 0 is the header
    return _ZEBRA_GRAY if data_row_idx % 2 == 1 else _ZEBRA_WHITE


def _section_header_row(table, label: str, n_cols: int):
    """Bold section label — mirrors frontend secHdr style."""
    row = table.add_row()
    _add_row_shading(row, _zebra_color(table))
    cell = row.cells[0]
    _set_cell_text(cell, label, bold=True, font_name="Arial", font_pt=10)
    for i in range(1, n_cols):
        cell.merge(row.cells[i])


def _data_row(table, label: str, *amounts, indent=True):
    """Indented label + right-aligned amounts — mirrors frontend label/num."""
    row = table.add_row()
    _add_row_shading(row, _zebra_color(table))
    prefix = "    " if indent else ""
    _set_cell_text(row.cells[0], prefix + label)
    for i, amt in enumerate(amounts, start=1):
        _set_cell_text(row.cells[i], amt, align_right=True)


def _total_row(table, label: str, *amounts, border_style="single"):
    """Bold total row with top border — mirrors frontend totCell/totNum.
    border_style='double' gives the double-underline grand-total style."""
    border_size = 12 if border_style == "double" else 8
    row = table.add_row()
    _add_row_shading(row, _zebra_color(table))
    _set_cell_text(row.cells[0], label, bold=True, font_name="Arial", font_pt=10)
    _add_cell_border(row.cells[0], "top", style=border_style, size=border_size)
    for i, amt in enumerate(amounts, start=1):
        _set_cell_text(row.cells[i], amt, bold=True, align_right=True)
        _add_cell_border(row.cells[i], "top", style=border_style, size=border_size)


def _header_row(table, col_headers: list[str]):
    """Column header row — bold, bottom border."""
    row = table.rows[0]
    for i, hdr in enumerate(col_headers):
        _set_cell_text(row.cells[i], hdr, bold=True,
                       align_right=(i > 0), font_name="Arial", font_pt=10)
        _add_cell_border(row.cells[i], "bottom", style="single", size=8)


def _make_table(doc: Document, n_cols: int):
    """Borderless table — borders added selectively via _add_cell_border."""
    table = doc.add_table(rows=1, cols=n_cols)
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        tblBorders.append(el)
    tblPr.append(tblBorders)
    return table


_COGS_NAMES                     = {"Cost of Goods Sold"}
_OPERATING_REVENUE_SUBTYPES     = {"Product Revenue", "Service Revenue"}
_OPERATING_EXPENSE_SUBTYPES     = {"Operating Expense"}
_NON_OPERATING_REVENUE_SUBTYPES = {"Non-Operating Revenue"}
_NON_OPERATING_EXPENSE_SUBTYPES = {"Non-Operating Expense"}
_TAX_SUBTYPES                   = {"Tax Expense"}


def _classify(r) -> str | None:
    """
    Section-membership rule for one income-statement row.

    MUST mirror classify() in IncomeStatementTable.jsx exactly — this is
    what the person sees on screen when reviewing/approving the statement,
    so the final docx/EDGAR HTML has to partition rows the same way.
    Checks acct_subtype first (e.g. "Interest Income" has subtype
    "Non-Operating Revenue" and belongs in Other Income, not Revenue),
    falling back to the coarser `category` field only if no subtype rule
    matches.
    """
    if getattr(r, "acct_name", "") in _COGS_NAMES:
        return "cogs"
    sub = getattr(r, "acct_subtype", None) or ""
    if sub in _OPERATING_REVENUE_SUBTYPES:
        return "op_revenue"
    if sub in _OPERATING_EXPENSE_SUBTYPES:
        return "op_expense"
    if sub in _NON_OPERATING_REVENUE_SUBTYPES:
        return "non_op_revenue"
    if sub in _NON_OPERATING_EXPENSE_SUBTYPES:
        return "non_op_expense"
    if sub in _TAX_SUBTYPES:
        return "tax"
    if getattr(r, "category", "") == "Revenue":
        return "op_revenue"
    if getattr(r, "category", "") == "Expense":
        return "op_expense"
    return None

# ── Cash Flow description sets (mirror CashFlowTable.jsx) ────────────────────

_CF_SECTION_DESCS = {
    "CASH FLOWS FROM OPERATING ACTIVITIES",
    "CASH FLOWS FROM INVESTING ACTIVITIES",
    "CASH FLOWS FROM FINANCING ACTIVITIES",
}
_CF_SUBTOTAL_DESCS = {
    "Net Cash from Operating Activities",
    "Net Cash from Investing Activities",
    "Net Cash from Financing Activities",
}
_CF_DOUBLE_DESCS = {"Net Increase in Cash", "Cash at End of Period"}


# ═══════════════════════════════════════════════════════════════════════════════
# Balance Sheet writer
# ═══════════════════════════════════════════════════════════════════════════════

def _write_balance_sheet(
    doc: Document,
    balance_sheet: list[BalanceSheetRow],
    col_label: str,
    prior_col: str,
):
    """Write Balance Sheet section — mirrors BalanceSheetTable.jsx."""
    doc.add_heading("Balance Sheets", level=3)

    assets      = [r for r in balance_sheet if r.category == "Asset"]
    liabilities = [r for r in balance_sheet if r.category == "Liability"]
    equity      = [r for r in balance_sheet if r.category == "Equity"]

    n_cols = 3
    table  = _make_table(doc, n_cols)
    _header_row(table, ["", col_label, prior_col])

    # Assets
    _section_header_row(table, "ASSETS", n_cols)
    for i, r in enumerate(assets):
        if i == 0:
            _data_row(table, r.acct_name, _add_dollar(_fmt(r.current_period)), _add_dollar(_fmt(r.prior_period)))
        else:
            _data_row(table, r.acct_name, _fmt(r.current_period), _fmt(r.prior_period))
    tot_a_cur = sum(r.current_period for r in assets)
    tot_a_pri = sum(r.prior_period   for r in assets)
    _total_row(table, "Total Assets", _fmt(tot_a_cur), _fmt(tot_a_pri))

    # Liabilities (stored positive credit-normal)
    _section_header_row(table, "LIABILITIES", n_cols)
    for r in liabilities:
        _data_row(table, r.acct_name,
                  _fmt(abs(r.current_period)), _fmt(abs(r.prior_period)))
    tot_l_cur = abs(sum(r.current_period for r in liabilities))
    tot_l_pri = abs(sum(r.prior_period   for r in liabilities))
    _total_row(table, "Total Liabilities", _fmt(tot_l_cur), _fmt(tot_l_pri))

    # Equity (pre-signed: positive additive, negative contra-equity)
    _section_header_row(table, "EQUITY", n_cols)
    for r in equity:
        _data_row(table, r.acct_name, _fmt(r.current_period), _fmt(r.prior_period))
    tot_e_cur = sum(r.current_period for r in equity)
    tot_e_pri = sum(r.prior_period   for r in equity)
    _total_row(table, "Total Equity", _fmt(tot_e_cur), _fmt(tot_e_pri))

    # Total L + E — double underline, last row of the statement -> $ sign
    _total_row(table, "Total Liabilities + Equity",
               _add_dollar(_fmt(tot_l_cur + tot_e_cur)),
               _add_dollar(_fmt(tot_l_pri + tot_e_pri)),
               border_style="double")
    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════════
# Income Statement writer
# ═══════════════════════════════════════════════════════════════════════════════

def _write_income_statement(
    doc: Document,
    income_stmt: list[IncomeStatementRow],
    col_label: str,
    prior_label: str,
    prior2_label: str = None,
    ytd_cur_label: str = None,
    ytd_prior_label: str = None,
):
    """
    Write Income Statement section — mirrors IncomeStatementTable.jsx's
    5-section multi-step GAAP layout:
        1. Revenue                                     -> Total Revenue
        2. Cost of Goods Sold                          -> Gross Profit
        3. Operating Expenses                          -> Operating Income / (Loss)
        4. Other Income / (Expenses)                   -> Income Before Income Tax
        5. Income Tax Expense                          -> NET INCOME / (LOSS)

    Rows are partitioned with _classify(), the same subtype-first rule the
    frontend uses — this is the fix for rows silently landing in the wrong
    section (e.g. "Interest Income" previously fell into a flat REVENUE
    bucket by category match; it belongs under Other Income / (Expenses)).
    Sections 2 and 4 are only emitted if there are rows in them, matching
    the frontend's conditional rendering.
    """
    doc.add_heading("Statements of Income", level=3)

    is_10k  = prior2_label is not None
    has_ytd = ytd_cur_label is not None

    if is_10k:
        n_cols  = 4 if prior2_label else 3
        headers = ["", col_label, prior_label]
        if prior2_label:
            headers.append(prior2_label)
    else:
        n_cols  = 5 if has_ytd else 3
        headers = ["", col_label, prior_label]
        if has_ytd:
            headers += [ytd_cur_label, ytd_prior_label]

    table = _make_table(doc, n_cols)
    _header_row(table, headers)

    def _get(r, field):
        return r.__dict__.get(field) or 0.0

    def _amounts(r, use_abs=False, negate=False):
        fn = _fmt_neg if negate else (_fmt_abs if use_abs else _fmt)
        vals = [fn(_get(r, "current_period")), fn(_get(r, "prior_period"))]
        if is_10k and prior2_label:
            vals.append(fn(_get(r, "prior2_period")))
        if not is_10k and has_ytd:
            vals += [fn(_get(r, "ytd_current")), fn(_get(r, "ytd_prior"))]
        return vals

    def _col_sum(rows, field):
        return sum(_get(r, field) for r in rows)

    def _sums(rows):
        """(cur, pri, p2_or_None, ytdc_or_None, ytdp_or_None) totals for a row group."""
        cur = _col_sum(rows, "current_period")
        pri = _col_sum(rows, "prior_period")
        p2   = _col_sum(rows, "prior2_period") if is_10k and prior2_label else None
        ytdc = _col_sum(rows, "ytd_current")   if not is_10k and has_ytd  else None
        ytdp = _col_sum(rows, "ytd_prior")     if not is_10k and has_ytd  else None
        return cur, pri, p2, ytdc, ytdp

    def _tot_vals(sums, use_abs=False):
        cur, pri, p2, ytdc, ytdp = sums
        fn = _fmt_abs if use_abs else _fmt
        vals = [fn(cur), fn(pri)]
        if is_10k and prior2_label:
            vals.append(fn(p2 or 0.0))
        if not is_10k and has_ytd:
            vals += [fn(ytdc or 0.0), fn(ytdp or 0.0)]
        return vals

    def _combine(a, b, op):
        """Elementwise combine two 5-tuples of (cur, pri, p2, ytdc, ytdp), skipping Nones."""
        return tuple(
            (op(x, y) if x is not None and y is not None else None)
            for x, y in zip(a, b)
        )

    op_rev_rows     = [r for r in income_stmt if _classify(r) == "op_revenue"]
    cogs_rows       = [r for r in income_stmt if _classify(r) == "cogs"]
    op_exp_rows     = [r for r in income_stmt if _classify(r) == "op_expense"]
    non_op_rev_rows = [r for r in income_stmt if _classify(r) == "non_op_revenue"]
    non_op_exp_rows = [r for r in income_stmt if _classify(r) == "non_op_expense"]
    tax_rows        = [r for r in income_stmt if _classify(r) == "tax"]

    # ── SECTION 1: Revenue ──
    _section_header_row(table, "REVENUE", n_cols)
    for i, r in enumerate(op_rev_rows):
        amts = _amounts(r)
        if i == 0:
            amts = [_add_dollar(a) for a in amts]
        _data_row(table, r.acct_name, *amts)
    rev_sums = _sums(op_rev_rows)
    _total_row(table, "Total Revenue", *_tot_vals(rev_sums))

    # ── SECTION 2: Cost of Goods Sold -> Gross Profit ──
    if cogs_rows:
        _section_header_row(table, "COST OF GOODS SOLD", n_cols)
        for r in cogs_rows:
            _data_row(table, r.acct_name, *_amounts(r, use_abs=True))
        cogs_sums = _sums(cogs_rows)
        gross_profit_sums = _combine(rev_sums, cogs_sums, lambda a, b: a - b)
        _total_row(table, "Gross Profit", *_tot_vals(gross_profit_sums))
    else:
        gross_profit_sums = rev_sums

    # ── SECTION 3: Operating Expenses -> Operating Income / (Loss) ──
    _section_header_row(table, "OPERATING EXPENSES", n_cols)
    for r in op_exp_rows:
        _data_row(table, r.acct_name, *_amounts(r, use_abs=True))
    op_exp_sums = _sums(op_exp_rows)
    _total_row(table, "Total Operating Expenses", *_tot_vals(op_exp_sums, use_abs=True))

    op_income_sums = _combine(gross_profit_sums, op_exp_sums, lambda a, b: a - b)
    _total_row(table, "Operating Income / (Loss)", *_tot_vals(op_income_sums))

    # ── SECTION 4: Other Income / (Expenses) -> Income Before Income Tax ──
    pretax_sums = op_income_sums
    if non_op_rev_rows or non_op_exp_rows:
        _section_header_row(table, "OTHER INCOME / (EXPENSES)", n_cols)
        for r in non_op_rev_rows:
            _data_row(table, r.acct_name, *_amounts(r))
        for r in non_op_exp_rows:
            # Parenthesized: these sit inside "Other Income / (Expenses)"
            # and are netted against non_op_rev_rows (e.g. Interest Income)
            # above -- the parens are the visual cue that this line
            # subtracts from that net total, matching the section label.
            _data_row(table, r.acct_name, *_amounts(r, negate=True))
        non_op_rev_sums = _sums(non_op_rev_rows)
        non_op_exp_sums = _sums(non_op_exp_rows)
        net_other_sums = _combine(non_op_rev_sums, non_op_exp_sums, lambda a, b: a - b)
        _total_row(table, "Total Other Income / (Expenses)", *_tot_vals(net_other_sums))
        pretax_sums = _combine(op_income_sums, net_other_sums, lambda a, b: a + b)

    _total_row(table, "Income Before Income Tax", *_tot_vals(pretax_sums))

    # ── SECTION 5: Income Tax Expense -> NET INCOME / (LOSS) ──
    tax_sums = (0.0, 0.0, 0.0 if is_10k and prior2_label else None,
                0.0 if not is_10k and has_ytd else None,
                0.0 if not is_10k and has_ytd else None)
    if tax_rows:
        for r in tax_rows:
            # No parens here -- unlike Interest Expense above, Income Tax
            # is its own standalone deduction line (Section 5), not netted
            # against another row in the same section.
            _data_row(table, _income_stmt_label(r), *_amounts(r, use_abs=True))
        tax_sums = _sums(tax_rows)

    # Net Income is computed as ONE direct pass over all revenue rows minus
    # all expense-like rows — the same shape financial_service.py's
    # net_income_for_period() uses (a single revenue-sum minus a single
    # expense-sum) — rather than derived algebraically through the chained
    # Gross Profit -> Operating Income -> Pretax Income -> Net Income
    # _combine() calls above. Those are algebraically equivalent but NOT
    # numerically guaranteed identical: floating-point +/- isn't
    # associative, so four chained operations can round to a different
    # dollar than one direct subtraction when the true value sits close to
    # a .5 boundary. That's what produced a real (if tiny) NetIncomeLoss
    # mismatch against the Cash Flow/Equity statements' Net Income, which
    # both call net_income_for_period() directly — same failure mode
    # net_income_for_period()'s own docstring documents fixing once already
    # for a similar FY2023 case. Filtering income_stmt (rather than
    # concatenating the already-split op_rev_rows/cogs_rows/etc. lists)
    # preserves the original chart-of-accounts iteration order, matching
    # net_income_for_period()'s own iteration order term-for-term.
    net_rev_rows = [r for r in income_stmt if _classify(r) in ("op_revenue", "non_op_revenue")]
    net_exp_rows = [r for r in income_stmt if _classify(r) in ("cogs", "op_expense", "non_op_expense", "tax")]
    net_rev_sums = _sums(net_rev_rows)
    net_exp_sums = _sums(net_exp_rows)
    net_income_sums = _combine(net_rev_sums, net_exp_sums, lambda a, b: a - b)
    net_income_vals = [_add_dollar(v) for v in _tot_vals(net_income_sums)]
    _total_row(table, "NET INCOME / (LOSS)", *net_income_vals, border_style="double")
    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════════
# Cash Flow Statement writer
# ═══════════════════════════════════════════════════════════════════════════════

def _write_cash_flow(
    doc: Document,
    cash_flow: list[CashFlowRow],
    col_label: str,
    prior_label: str,
    prior2_label: str = None,
):
    """Write Cash Flow Statement section — mirrors CashFlowTable.jsx."""
    doc.add_heading("Statements of Cash Flows", level=3)

    show_p2 = prior2_label is not None
    n_cols  = 4 if show_p2 else 3
    headers = ["", col_label, prior_label]
    if show_p2:
        headers.append(prior2_label)

    table = _make_table(doc, n_cols)
    _header_row(table, headers)

    first_data_row_seen = False
    last_idx = len(cash_flow) - 1
    for i, r in enumerate(cash_flow):
        desc = r.description or ""
        amts = [_fmt(r.current_period), _fmt(r.prior_period)]
        if show_p2:
            amts.append(_fmt(r.prior2_period))

        if desc in _CF_SECTION_DESCS:
            _section_header_row(table, desc, n_cols)
            continue

        is_dollar_row = (not first_data_row_seen) or (i == last_idx)
        if is_dollar_row:
            amts = [_add_dollar(a) for a in amts]
        first_data_row_seen = True

        if desc in _CF_SUBTOTAL_DESCS:
            _total_row(table, desc, *amts, border_style="single")
        elif desc in _CF_DOUBLE_DESCS:
            _total_row(table, desc, *amts, border_style="double")
        else:
            _data_row(table, desc, *amts)

    doc.add_paragraph()


_EQUITY_SECTION_DESC_RE = re.compile(r"^(?:Three|Six|Nine) Months Ended$")


def _is_equity_section_header(description: str) -> bool:
    """
    True for a stockholders'-equity rollforward section-header row
    ("Three Months Ended" / "Six Months Ended" / "Nine Months Ended").
    Pattern-based rather than a fixed set of exact strings, so a new
    duration label (e.g. Q3's "Nine Months Ended") is recognized
    automatically as soon as financial_service.py starts producing it,
    without this file also needing an update in lockstep.
    """
    return bool(_EQUITY_SECTION_DESC_RE.match(description))


def _write_stockholders_equity(doc: Document, equity_rows: list) -> None:
    """
    Write Statement of Stockholders' Equity — mirrors
    StockholdersEquityTable.jsx exactly: a matrix table (one column per
    equity component, one row per balance/activity line), rather than the
    period-column layout the other three statements use.

    Total Equity is bold on EVERY row, not just balance rows — matches the
    reference grid (equity-grid-gem.docx), where e.g. "Net Income" shows a
    bold total even though the row itself isn't a balance row. Balance
    rows get the same bold + top-border treatment _total_row() gives
    financial-statement subtotals; the LAST balance row gets the double
    border, matching the grand-total convention everywhere else in this
    file.

    Two additional row kinds exist for the 10-Q's 4-block layout (see
    build_stockholders_equity_quarterly() in financial_service.py) and
    never appear in the 10-K's continuous 3-year chain:
      - A blank separator row (description == "") between each pair of
        independent rollforward blocks — rendered as a fully empty row,
        not five "—" placeholders.
      - A section-header row (description in _EQUITY_SECTION_DESCS, i.e.
        "Three Months Ended" / "Six Months Ended") — rendered as a bold
        label spanning the row, same convention _section_header_row() uses
        for the Cash Flow statement's own section dividers, rather than
        trying to format five empty numeric columns.
    """
    doc.add_heading("Statements of Stockholders' Equity", level=3)

    headers = [
        "", "Common Stock", "Treasury Stock", "Retained Earnings",
        "Accumulated Other\n Comprehensive Income (Loss)",
        "Total Stockholders'\n Equity",
    ]
    n_cols = len(headers)
    table = _make_table(doc, n_cols)
    _header_row(table, headers)

    balance_indices = [i for i, r in enumerate(equity_rows) if r.is_balance_row]
    last_balance_idx = balance_indices[-1] if balance_indices else None

    for i, r in enumerate(equity_rows):
        if _is_equity_section_header(r.description):
            _section_header_row(table, r.description, n_cols)
            continue
        if not r.description and not r.is_balance_row:
            # Blank separator row between independent rollforward blocks —
            # a genuinely empty row, not five "—" placeholders.
            table.add_row()
            continue

        if r.is_balance_row:
            amts = [
                _add_dollar(_fmt(r.common_stock_amount)),
                _add_dollar(_fmt(r.treasury_stock)),
                _add_dollar(_fmt(r.retained_earnings)),
                _add_dollar(_fmt(r.accumulated_oci)),
                _add_dollar(_fmt(r.total_equity)),
            ]
            border_style = "double" if i == last_balance_idx else "single"
            _total_row(table, r.description, *amts, border_style=border_style)
        else:
            amts = [
                _fmt(r.common_stock_amount),
                _fmt(r.treasury_stock),
                _fmt(r.retained_earnings),
                _fmt(r.accumulated_oci),
                _fmt(r.total_equity),
            ]
            row = table.add_row()
            _add_row_shading(row, _zebra_color(table))
            _set_cell_text(row.cells[0], r.description)
            last_col = len(amts)
            for j, amt in enumerate(amts, start=1):
                _set_cell_text(row.cells[j], amt, bold=(j == last_col), align_right=True)

    doc.add_paragraph()


# ─── 10-Q Item 1: Financial Statements ───────────────────────────────────────

def create_10q_item010(
    balance_sheet: list[BalanceSheetRow],
    income_stmt: list[IncomeStatementRow],
    cash_flow: list[CashFlowRow],
    stockholders_equity: list,
    period_label: str,
    prior_label: str,
    bs_prior_label: str,
    year: int,
    quarter: int,
) -> str:
    """Create item010-Financial-Statements.docx for 10-Q Part I Item 1.
    Writes to DATA_10Q_PART1 (container path), which is bind-mounted
    to WIN_10Q_PART1 on the Windows host.

    bs_prior_label  — balance sheet prior column header (prior fiscal year-end),
                      e.g. "December 31, 2024".  Matches the frontend
                      BalanceSheetTable priorCol logic.
    """
    _ensure_dir(os.path.join(_q10_base(quarter), settings.DATA_10Q_PART1))
    output_path = os.path.join(_q10_base(quarter), settings.DATA_10Q_PART1, "item010-Financial-Statements.docx")
    doc = Document()

    _add_heading(doc, "Item 1. Financial Statements", level=3)

    # Determine YTD column headers (Q2/Q3 only)
    YTD_DURATION = {1: "Three Months", 2: "Six Months", 3: "Nine Months"}
    dur = YTD_DURATION.get(quarter, "")
    # build_income_statement() always computes ytd_current/ytd_prior for
    # every quarter, including Q1 - Q1's own YTD start/end happen to equal
    # the quarter's own start/end (the first quarter's YTD *is* the
    # quarter), so ytd_current is never None even in Q1. Without excluding
    # quarter 1 explicitly here, that produced two redundant "Three Months
    # Ended" column pairs side by side instead of omitting the YTD columns
    # entirely, as a real 10-Q does for Q1.
    has_ytd = quarter != 1 and income_stmt and income_stmt[0].ytd_current is not None
    # Case-insensitive split: period_label/prior_label are built with title
    # case ("Three Months Ended June 30, 2025"), but this used to split on
    # the lowercase literal 'ended', which never matched — so .split()
    # silently returned the ENTIRE original label unchanged, and it got
    # embedded whole into the new label (e.g. "Six Months ended Three
    # Months Ended June 30, 2025"). re.split with IGNORECASE actually
    # strips the "<duration> Ended " prefix, leaving just the date.
    # Case-insensitive split: period_label/prior_label are built with title
    # case ("Three Months Ended\n June 30, 2025") — a literal "\n " before
    # the date is what makes the header wrap onto a second line, matching
    # the 10-K's "Year ended\n December 31, 2025" convention. .strip()ping
    # the split result removes that newline entirely, so the YTD columns
    # built here need it re-added explicitly rather than reusing whatever
    # whitespace .strip() left behind.
    ytd_cur_label   = f"{dur} ended\n {re.split('ended', period_label, flags=re.IGNORECASE)[-1].strip()}" if has_ytd else None
    ytd_prior_label = f"{dur} ended\n {re.split('ended', prior_label,  flags=re.IGNORECASE)[-1].strip()}" if has_ytd else None

    # Cash flow always needs a real label, every quarter including Q1 -
    # it's inherently presented YTD (see Report-10-Q-design.docx: "Cash
    # Flow Statement: for the year-to-date period..."), unlike the income
    # statement's EXTRA columns above, which only apply to Q2/Q3 and are
    # correctly None for Q1. Reusing ytd_cur_label/ytd_prior_label here
    # meant Q1's cash flow headers went None/None the moment has_ytd
    # became False for Q1 - they need their own always-computed pair
    # instead of sharing the income statement's has_ytd gate.
    cash_flow_cur_label   = f"{dur} ended\n {re.split('ended', period_label, flags=re.IGNORECASE)[-1].strip()}"
    cash_flow_prior_label = f"{dur} ended\n {re.split('ended', prior_label,  flags=re.IGNORECASE)[-1].strip()}"

    # Balance sheet uses bs_prior_label (prior fiscal year-end date),
    # not the quarter-comparison prior_label used by income / cash flow.
    _write_balance_sheet(doc, balance_sheet, period_label, bs_prior_label or prior_label)
    _write_income_statement(
        doc, income_stmt, period_label, prior_label,
        ytd_cur_label=ytd_cur_label,
        ytd_prior_label=ytd_prior_label,
    )
    _write_cash_flow(doc, cash_flow, cash_flow_cur_label, cash_flow_prior_label)
    if stockholders_equity:
        _write_stockholders_equity(doc, stockholders_equity)

    doc.save(output_path)
    return output_path


# ─── 10-Q Item 2: MD&A ────────────────────────────────────────────────────────

def create_10q_item020(narrative: str, quarter: int) -> str:
    """Create item020-MDA.docx for 10-Q Part I Item 2.
    Writes to the quarterly DATA_10Q_PART1 (container path) — see
    _q10_base()'s docstring.
    """
    _ensure_dir(os.path.join(_q10_base(quarter), settings.DATA_10Q_PART1))
    output_path = os.path.join(_q10_base(quarter), settings.DATA_10Q_PART1, "item020-MDA.docx")
    doc = Document()
    _add_heading(doc, "Item 2. Management's Discussion and Analysis", level=3)
    parser = HtmlToDocx()
    parser.add_html_to_document(narrative, doc)
    doc.save(output_path)
    return output_path


# ─── 10-Q Item 3 (item011): Notes to Financial Statements ───────────────────

def create_10q_item011(narrative: str, quarter: int) -> str:
    """Create item011-Notes-to-Financial-Statements.docx for 10-Q Part I Item 1 (Notes).
    Writes to the quarterly DATA_10Q_PART1 (container path) — see
    _q10_base()'s docstring. The narrative is an HTML string produced by
    the TipTap editor.
    """
    _ensure_dir(os.path.join(_q10_base(quarter), settings.DATA_10Q_PART1))
    output_path = os.path.join(
        _q10_base(quarter), settings.DATA_10Q_PART1, "item011-Notes-to-Financial-Statements.docx"
    )
    doc = Document()
    _add_heading(
        doc,
        "Notes to Condensed Consolidated Financial Statements (Unaudited)",
        level=3,
    )
    parser = HtmlToDocx()
    parser.add_html_to_document(narrative, doc)
    _right_align_notes_table_columns(doc)
    doc.save(output_path)
    return output_path


# ─── 10-Q Final Report ────────────────────────────────────────────────────────

def _merge_dir_into_composer(composer: Composer, directory: str, heading: str = None):
    """Merge all .docx files from a directory (sorted) into the composer."""
    if heading:
        p = composer.doc.add_paragraph()
        run = p.add_run(heading)
        run.bold = True
        p.style = composer.doc.styles["Heading 3"]
    if not os.path.exists(directory):
        return
    for f in sorted(Path(directory).glob("*.docx")):
        composer.append(Document(str(f)))


def create_10q_final_report(year: int, quarter: int) -> str:
    """Merge all 10-Q item files into sec-10q-{year}-q{quarter}.docx.
    Writes to REPORTS_DIR. Year/quarter are suffixed onto the filename so
    each period's report is kept as its own file rather than overwriting
    whatever the last run produced.
    """
    _ensure_dir(settings.REPORTS_DIR)
    output_path = os.path.join(settings.REPORTS_DIR, f"sec-10q-{year}-q{quarter}.docx")
    # logger.info(f"before merge, REPORTS output_path: {output_path}")
    master = Document()
    composer = Composer(master, preserve_styles=True)

    _merge_dir_into_composer(composer, os.path.join(_q10_base(quarter), settings.DATA_10Q_INTRO))
    _merge_dir_into_composer(composer, os.path.join(_q10_base(quarter), settings.DATA_10Q_PART1), heading="PART I.")
    _merge_dir_into_composer(composer, os.path.join(_q10_base(quarter), settings.DATA_10Q_PART2), heading="PART II.")
    master.add_page_break()
    _merge_dir_into_composer(composer, os.path.join(_q10_base(quarter), settings.DATA_10Q_SIGS))
    master.save(output_path)
    return output_path


def _merge_10q_body_only_docx(quarter: int) -> str:
    """
    Merge Part I-II + Signatures into a temp docx — same as
    create_10q_final_report(), but WITHOUT the intro directory, since the
    intro (cover + TOC) is converted separately by convert_intro_to_html()
    and spliced in as styled HTML instead. Mirrors _merge_10k_body_only_docx()
    exactly, just with the 10-Q's own two-part (not four-part) structure,
    and reading from the quarterly directory (see _q10_base()'s docstring).
    """
    _ensure_dir(settings.REPORTS_DIR)
    tmp_path = os.path.join(settings.REPORTS_DIR, "_sec-10q-body-only.docx")

    master = Document()
    composer = Composer(master, preserve_styles=True)
    _merge_dir_into_composer(composer, os.path.join(_q10_base(quarter), settings.DATA_10Q_PART1), heading="PART I.")
    _merge_dir_into_composer(composer, os.path.join(_q10_base(quarter), settings.DATA_10Q_PART2), heading="PART II.")
    master.add_page_break()
    _merge_dir_into_composer(composer, os.path.join(_q10_base(quarter), settings.DATA_10Q_SIGS))

    master.save(tmp_path)
    return tmp_path


def create_10q_edgar_html(output_filename: str, quarter: int) -> str:
    """
    Build the final EDGAR HTML for the 10-Q by merging TWO separately
    generated pieces, BOTH going through the SAME styled walker
    (_docx_body_to_html_parts(preserve_style=True)) used by
    create_10k_edgar_html() — see that function's docstring for the full
    rationale (EDGAR strips <style> blocks per EFM 5.02.05, so inline
    style="..." attributes are the only presentational info that survives):
      1. item001 (cover) + item002 (TOC), via convert_intro_to_html().
      2. Part I + Part II + Signatures (financial statements, MD&A, Notes,
         etc.).

    This replaces the prior 10-Q behavior of falling back to
    convert_to_edgar_html() (the bare, unstyled legacy path) — the 10-Q now
    gets the identical styled-paragraph / real-inline-style / <strong>-wrapped
    tables the 10-K already has, instead of losing all formatting once
    EDGAR strips the <style> block.

    xbrl_tagger.py's table/notes tagging is unaffected by this switch, for
    the same reason noted in create_10k_edgar_html(): _tag_table /
    tag_notes_section read cell/paragraph text via BeautifulSoup's
    .get_text(), which is blind to inline style attributes and <strong>
    wrappers; only the raw-string regex passes (tag_cover_page,
    notes_config.py's heading/end patterns) care about markup shape, and
    are already tolerant of an optional style="..." attribute and an
    optional <strong> wrapper.

    Call fill_cover_page(period_end) before this so item001 reflects the
    current business_info.toml. xbrl_tagger.tag_filing() should run on the
    output of THIS function for the final 10-Q — not convert_to_edgar_html().
    """
    _ensure_dir(settings.REPORTS_DIR)
    output_path = os.path.join(settings.REPORTS_DIR, output_filename)

    intro_paths = sorted(Path(os.path.join(_q10_base(quarter), settings.DATA_10Q_INTRO)).glob("*.docx"))
    intro_html = convert_intro_to_html([str(p) for p in intro_paths])

    body_docx_path = _merge_10q_body_only_docx(quarter)
    body_doc = Document(body_docx_path)
    body_parts = _docx_body_to_html_parts(body_doc, preserve_style=True)
    # Same h3-realignment fix as the 10-K: _docx_body_to_html_parts hardcodes
    # text-align: center on every heading level, correct only for the
    # intro's Heading 1/2. Every Heading 3 in this body (Part headings, Item
    # headings, "Notes to Condensed Consolidated Financial Statements", each
    # note's own <h3>, etc.) is left-aligned in Ron's source docx files.
    body_html = _reset_h3_align_left("\n".join(body_parts))

    # TOC <-> section jump links: needs the intro (TOC) and body (section
    # headings) together in one pass, but must stay OUTSIDE the
    # DOCTYPE/meta/<style> wrapper below — see
    # _add_toc_navigation_links()'s docstring.
    linked_intro_and_body = _add_toc_navigation_links(intro_html + "\n" + body_html)

    html_parts = [
        "<?xml version='1.0' encoding='UTF-8'?>",
        "<!DOCTYPE html PUBLIC \"-//W3C//DTD XHTML 1.0 Strict//EN\" "
        "\"http://www.w3.org/TR/xhtml1/DTD/xhtml1-strict.dtd\">",
        "<html xmlns='http://www.w3.org/1999/xhtml'>",
        "<head>",
        "<meta http-equiv='Content-Type' content='text/html; charset=UTF-8'/>",
        f"<title>{output_filename.replace('.htm', '')}</title>",
        "<style type='text/css'>", _EDGAR_CSS, "</style>",
        "</head>",
        "<body>",
        linked_intro_and_body,
    ]
    html_parts.extend(["</body>", "</html>"])

    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(html_parts))

    try:
        os.remove(body_docx_path)
    except OSError:
        pass

    return output_path

def create_10k_item081(
    balance_sheet: list[BalanceSheetRow],
    income_stmt: list[IncomeStatementRow],
    cash_flow: list[CashFlowRow],
    stockholders_equity: list[StockholdersEquityRow],
    period_label: str,
    prior_label: str,
    prior2_label: str,
    year: int,
) -> str:
    """Create item081-Financial-Statements.docx for 10-K Item 8.
    Writes to DATA_10K_PART2 (container path).
    """
    _ensure_dir(settings.DATA_10K_PART2)
    output_path = os.path.join(settings.DATA_10K_PART2, "item081-Financial-Statements.docx")
    doc = Document()

    _add_heading(doc, "Financial Statements", level=3)

    has_prior2 = income_stmt and income_stmt[0].prior2_period is not None
    p2_lbl = prior2_label if has_prior2 else None

    _write_balance_sheet(doc, balance_sheet, period_label, prior_label)
    _write_income_statement(doc, income_stmt, period_label, prior_label,
                            prior2_label=p2_lbl)
    _write_cash_flow(doc, cash_flow, period_label, prior_label,
                     prior2_label=p2_lbl)
    if stockholders_equity:
        _write_stockholders_equity(doc, stockholders_equity)

    doc.save(output_path)
    return output_path


# ─── 10-K Item 7: MD&A ────────────────────────────────────────────────────────

def create_10k_item070(narrative: str) -> str:
    """Create item070-MDA.docx for 10-K Item 7.
    Writes to DATA_10K_PART2 (container path), bind-mounted to WIN_10K_PART2.
    """
    _ensure_dir(settings.DATA_10K_PART2)
    output_path = os.path.join(settings.DATA_10K_PART2, "item070-MDA.docx")
    doc = Document()
    _add_heading(
        doc,
        "Item 7. Management's Discussion and Analysis of Financial Condition and Results of Operations",
        level=3,
    )
    parser = HtmlToDocx()
    parser.add_html_to_document(narrative, doc)
    doc.save(output_path)
    return output_path


# ─── 10-K Item 8 (item082): Notes to Financial Statements ───────────────────

def create_10k_item082(narrative: str) -> str:
    """Create item082-Notes-to-Financial-Statements.docx for 10-K Item 8.
    Writes to DATA_10K_PART2 (container path), bind-mounted to WIN_10K_PART2.
    The narrative is an HTML string produced by the TipTap editor.
    """
    _ensure_dir(settings.DATA_10K_PART2)
    output_path = os.path.join(
        settings.DATA_10K_PART2, "item082-Notes-to-Financial-Statements.docx"
    )
    doc = Document()
    _add_heading(
        doc,
        "Notes to Consolidated Financial Statements",
        level=3,
    )
    parser = HtmlToDocx()
    parser.add_html_to_document(narrative, doc)
    _right_align_notes_table_columns(doc)
    doc.save(output_path)
    return output_path


# ─── Styled HTML preview (human-readable only — NOT used for iXBRL tagging) ──

def convert_to_styled_preview_html(docx_path: str, title: str = "Document Preview") -> str:
    """
    Render a docx into richly-styled HTML (centering, bold, table borders,
    per-paragraph font/spacing all preserved) using docx-parser-converter.

    IMPORTANT — this is a SEPARATE code path from convert_to_edgar_html().
    Do NOT feed this function's output into xbrl_tagger.py. Every regex in
    tag_cover_page(), tag_notes_section(), and notes_config.py's
    heading_pattern/end_pattern was written against convert_to_edgar_html()'s
    plain, un-styled markup (bare <p>, <th> for single-row tables, no
    surrounding <strong>/inline-style wrappers). docx-parser-converter's
    output is structurally different (inline style="..." on every element,
    <strong> wrapping, <td> instead of <th> for single-row tables, etc.) and
    would silently break those regex matches if it ever replaced
    convert_to_edgar_html() in the tagging path.

    Use this only for human-facing previews — e.g. an endpoint that lets the
    frontend show what the cover page / a merged report actually looks like
    before or instead of the plain EDGAR-tagged version.

    Known caveat: docx-parser-converter can inject a stray
    <span class="list-marker">N</span> before a Heading1-styled paragraph
    that Word's numbering.xml treats as part of an outline (observed on
    "FORM 10-K" in item001-10-k-cover.docx). Harmless for a preview, but
    worth knowing if you see an unexpected "1" before a heading.
    """
    from docx_parser_converter import docx_to_html, ConversionConfig

    config = ConversionConfig(
        title=title,
        style_mode="inline",
        use_semantic_tags=True,
        table_mode="auto",
        fragment_only=True,
    )
    return docx_to_html(docx_path, config=config)


# ─── 10-K Cover Page (DEI-driven, via docxtpl) ───────────────────────────────

def _load_business_info() -> dict:
    """Read business_info.toml — same file/location xbrl_tagger.py and
    notes_agent.py read. Returns {} if missing/unparsable."""
    toml_path = os.path.join(settings.DATA_USER_INPUT_DIR, "business_info.toml")
    try:
        with open(toml_path, "rb") as f:
            return tomllib.load(f)
    except FileNotFoundError:
        logger.warning(f"business_info.toml not found at {toml_path}; cover page placeholders will be blank.")
        return {}
    except Exception as exc:
        logger.warning(f"Could not parse business_info.toml: {exc}; cover page placeholders will be blank.")
        return {}


def _format_fiscal_year_end_date(period_end: str) -> str:
    """'2025-12-31' -> 'December 31, 2025' (spelled-out month, matches the
    ixt:date-monthname-day-year-en transform xbrl_tagger.py expects)."""
    from datetime import date as _date
    d = _date.fromisoformat(period_end)
    return f"{d.strftime('%B')} {d.day}, {d.year}"


def fill_cover_page(period_end: str, template_path: str = None) -> str:
    """
    Render item001-10-k-cover.docx from its docxtpl placeholder template,
    substituting business_info.toml values (entity name, EIN, address,
    filer status, etc.) plus the filing's own fiscal-year-end date.

    Writes the rendered docx into settings.DATA_10K_INTRO — the same
    directory create_10k_final_report()'s docxcompose merge reads from —
    overwriting any previous rendering. The PRISTINE template (with
    {{ placeholders }} still intact) lives in settings.DATA_USER_INPUT_DIR
    (the same folder as business_info.toml) by default, so it's never in
    the same place as — and never gets clobbered by — the rendered output.
    Override via settings.COVER_TEMPLATE_PATH or the template_path arg if
    you keep it somewhere else.

    Call this once per report generation, before create_10k_final_report().
    """
    from docxtpl import DocxTemplate

    template_path = template_path or getattr(
        settings, "COVER_TEMPLATE_PATH",
        os.path.join(settings.DATA_USER_INPUT_DIR, "item001-10-k-cover.docx"),
    )
    biz = _load_business_info()

    context = {
        "entity_name":             biz.get("company_name", "[Company Name]"),
        "commission_file_number":  biz.get("commission_file_number", ""),
        "state_of_incorporation":  biz.get("state_of_incorporation", ""),
        "ein":                     biz.get("ein", ""),
        "address_line1":           biz.get("address_line1", ""),
        "city":                    biz.get("city", ""),
        "state_abbr":              biz.get("state_abbr", ""),
        "zip_code":                biz.get("zip_code", ""),
        "area_code":               biz.get("area_code", ""),
        "local_phone_number":      biz.get("local_phone_number", ""),
        "security_class":          biz.get("security_class", ""),
        "ticker":                  biz.get("ticker", ""),
        "exchange":                biz.get("exchange", ""),
        "aggregate_market_value":  biz.get("aggregate_market_value", "0"),
        "shares_outstanding":      biz.get("shares_outstanding", "0"),
        "fiscal_year_end_date":    _format_fiscal_year_end_date(period_end),
        "filer_category":                biz.get("filer_category", "non_accelerated"),
        "is_smaller_reporting_company":  biz.get("is_smaller_reporting_company", True),
        "is_emerging_growth_company":    biz.get("is_emerging_growth_company", False),
        "is_well_known_seasoned_issuer": biz.get("is_well_known_seasoned_issuer", False),
        "is_exempt_from_filing":         biz.get("is_exempt_from_filing", False),
        "filed_all_required_reports":    biz.get("filed_all_required_reports", True),
        "submitted_interactive_data":    biz.get("submitted_interactive_data", True),
    }

    tpl = DocxTemplate(template_path)
    tpl.render(context)

    _ensure_dir(settings.DATA_10K_INTRO)
    output_path = os.path.join(settings.DATA_10K_INTRO, "item001-10-k-cover.docx")
    tpl.save(output_path)
    # logger.info(f"[fill_cover_page] rendered cover page -> {output_path}")
    return output_path


# ─── 10-Q Cover Page (DEI-driven, via docxtpl) ────────────────────────────────
# Mirrors fill_cover_page() (10-K) above. item001-10-q-cover.docx already
# uses {{ quarter_end_date }} correctly in both places it appears ("For
# the quarterly period ended ..." and "As of ..., N shares ... issued and
# outstanding.") — the bug was never a missing/static template, it was
# report_10q.py calling fill_cover_page() (the 10-K function, which
# renders item001-10-k-cover.docx into DATA_10K_INTRO) instead of this
# function, so item001-10-q-cover.docx was never actually re-rendered
# per quarter.

def fill_cover_page_10q(period_end: str, quarter: int, template_path: str = None) -> str:
    """
    Render item001-10-q-cover.docx from its docxtpl placeholder template,
    substituting business_info.toml values plus THIS quarter's own period
    end date.

    Writes into _q10_base(quarter)/DATA_10Q_INTRO — the same directory
    create_10q_final_report()'s docxcompose merge AND
    create_10q_edgar_html()'s convert_intro_to_html() read
    item001-10-q-cover.docx from — overwriting any previous rendering.

    Call this once per report generation, before create_10q_final_report()
    and before create_10q_edgar_html(), passing the actual quarter-end
    date for the year/quarter being filed (e.g. from _quarter_dates()).
    """
    from docxtpl import DocxTemplate

    template_path = template_path or getattr(
        settings, "COVER_TEMPLATE_PATH_10Q",
        os.path.join(settings.DATA_USER_INPUT_DIR, "item001-10-q-cover.docx"),
    )
    biz = _load_business_info()

    context = {
        "entity_name":             biz.get("company_name", "[Company Name]"),
        "commission_file_number":  biz.get("commission_file_number", ""),
        "state_of_incorporation":  biz.get("state_of_incorporation", ""),
        "ein":                     biz.get("ein", ""),
        "address_line1":           biz.get("address_line1", ""),
        "city":                    biz.get("city", ""),
        "state_abbr":              biz.get("state_abbr", ""),
        "zip_code":                biz.get("zip_code", ""),
        "area_code":               biz.get("area_code", ""),
        "local_phone_number":      biz.get("local_phone_number", ""),
        "security_class":          biz.get("security_class", ""),
        "ticker":                  biz.get("ticker", ""),
        "exchange":                biz.get("exchange", ""),
        "shares_outstanding":      biz.get("shares_outstanding", "0"),
        # Matches the template's own placeholder name exactly — the 10-K's
        # fill_cover_page() calls this "fiscal_year_end_date"; the 10-Q
        # template calls it "quarter_end_date" since it's a quarter end,
        # not a fiscal year end. Drives both the "quarterly period ended"
        # sentence and the shares-outstanding sentence.
        "quarter_end_date":              _format_fiscal_year_end_date(period_end),
        "filer_category":                biz.get("filer_category", "non_accelerated"),
        "is_smaller_reporting_company":  biz.get("is_smaller_reporting_company", True),
        "is_emerging_growth_company":    biz.get("is_emerging_growth_company", False),
        "filed_all_required_reports":    biz.get("filed_all_required_reports", True),
        "submitted_interactive_data":    biz.get("submitted_interactive_data", True),
    }

    tpl = DocxTemplate(template_path)
    tpl.render(context)

    intro_dir = os.path.join(_q10_base(quarter), settings.DATA_10Q_INTRO)
    _ensure_dir(intro_dir)
    output_path = os.path.join(intro_dir, "item001-10-q-cover.docx")
    tpl.save(output_path)
    # logger.info(f"[fill_cover_page_10q] rendered cover page -> {output_path}")
    return output_path


# ─── 10-K Final Report ────────────────────────────────────────────────────────

def create_10k_final_report(year: int) -> str:
    """Merge all 10-K item files into sec-10k-{year}.docx.
    Writes to REPORTS_DIR (container path), bind-mounted to WIN_REPORTS_DIR.
    Year is suffixed onto the filename so each fiscal year's report is kept
    as its own file rather than overwriting whatever the last run produced.

    Call fill_cover_page(period_end) before this, so item001-10-k-cover.docx
    in DATA_10K_INTRO reflects the current business_info.toml before it gets
    merged in below.
    """
    _ensure_dir(settings.REPORTS_DIR)
    output_path = os.path.join(settings.REPORTS_DIR, f"sec-10k-{year}.docx")

    master = Document()
    composer = Composer(master, preserve_styles=True)

    _merge_dir_into_composer(composer, settings.DATA_10K_INTRO)
    _merge_dir_into_composer(composer, settings.DATA_10K_PART1, heading="PART I.")
    _merge_dir_into_composer(composer, settings.DATA_10K_PART2, heading="PART II.")
    _merge_dir_into_composer(composer, settings.DATA_10K_PART3, heading="PART III.")
    _merge_dir_into_composer(composer, settings.DATA_10K_PART4, heading="PART IV.")
    master.add_page_break()
    _merge_dir_into_composer(composer, settings.DATA_10K_SIGS)

    master.save(output_path)
    return output_path


# ─── EDGAR HTML Conversion ────────────────────────────────────────────────────

# Shared CSS for every EDGAR HTML output (cover-fragment splice included).
# Only additions here are BLANKET, tag-type-level rules (never per-element
# inline attributes) — so this is safe to apply everywhere: it changes how
# things look, never the tag/attribute shape that xbrl_tagger.py's regexes
# match against.
_EDGAR_CSS = """
body { font-family: Times New Roman, serif; font-size: 10pt; }
h1 { font-size: 14pt; font-weight: bold; text-align: center; }
h2 { font-size: 12pt; font-weight: bold; text-align: center; }
h3 { font-size: 11pt; font-weight: bold; text-align: center; }
table { border-collapse: collapse; width: 100%; margin: 10px 0; }
td { border: 1px solid #000; padding: 4px 8px; }
th { border: 1px solid #000; padding: 4px 8px; text-align: center; }
th { background-color: #ffffff; font-weight: bold; }
.number { text-align: right; }
"""


def _docx_body_to_html_parts(doc: Document, preserve_style: bool = False) -> list:
    """
    Shared DOCX -> HTML body walker.

    preserve_style=False: EXACT original plain behavior — bare <p>,
    <th>/<td>, <h1>/<h2>/<h3>, no alignment/bold preserved, no font
    declarations. Used only by convert_to_edgar_html(), the legacy plain
    conversion path (kept for reference / non-EDGAR previews — NOT used
    for the tagged final 10-K anymore).

    preserve_style=True: emits a real inline style="..." attribute on
    every paragraph (body font/size always; text-align only when the
    source docx explicitly centers/right-aligns/justifies it), a real
    font+centering style on every heading, an inline style on every
    table/cell (borders, alignment, width for item-listing tables), and
    wraps bold runs in <strong>. This is now used for BOTH pieces of the
    final 10-K EDGAR HTML — convert_intro_to_html() (item001/item002) AND
    create_10k_edgar_html()'s own body walk (Parts I-IV + Signatures,
    including the financial statements and Notes to Financial
    Statements) — since EDGAR strips <style> blocks entirely (EFM
    5.02.05) and inline style="..." is the only presentational
    information that survives into the filed .htm.

    Every raw-string regex in xbrl_tagger.py that runs against this
    preserve_style=True output BEFORE BeautifulSoup parses the document
    (tag_cover_page, tag_auditor_report_block, tag_notes_section) — plus
    notes_config.py's/notes_registry.py's heading/end patterns — is
    written to tolerate exactly what this adds: an optional attribute on
    the opening tag (e.g. a bare `<p>` becomes tolerant of `<p ...>` in
    the regex) and an optional `<strong>...</strong>` wrapper around the
    matched text. xbrl_tagger.py's financial-statement table tagging
    (_tag_table / _tag_stockholders_equity_table) needs no such
    tolerance at all — it walks the already-parsed BeautifulSoup tree
    and reads cell text via .get_text(), which strips every tag (inline
    style attributes and <strong> wrappers included) before the text is
    ever compared against GAAP_MAPPINGS or parsed as a number.
    """
    from docx.oxml.ns import qn

    _JC_MAP = {"center": "center", "end": "right", "right": "right", "both": "justify"}

    def _para_jc(p_element) -> str:
        """Raw w:jc value ('center'/'end'/'right'/'both') or '' if unset."""
        ppr = p_element.find(qn('w:pPr'))
        if ppr is None:
            return ""
        jc = ppr.find(qn('w:jc'))
        if jc is None:
            return ""
        return jc.get(qn('w:val'), "")

    def _para_align_style(p_element) -> str:
        if not preserve_style:
            return ""
        align = _JC_MAP.get(_para_jc(p_element), "")
        return f' style="text-align: {align}"' if align else ""

    def _cell_jc(cell_elem) -> str:
        """First paragraph in the cell that has an explicit w:jc, if any."""
        for para in cell_elem.findall('.//' + qn('w:p')):
            jc = _para_jc(para)
            if jc:
                return jc
        return ""

    def _run_inner_html(run) -> str:
        """
        Convert a single <w:r> run's children to an HTML string, walking
        them in document order. A run can contain more than one <w:t>
        separated by a manual <w:br/> (e.g. a hard-wrapped column header
        like "Accumulated Other"<w:br/>"Comprehensive Income (Loss)"), and
        that ordering/interleaving is lost if you just findall() every
        <w:t> descendant and join their text — the <w:br/> disappears
        entirely, silently concatenating the two text fragments with no
        space and no line break. Walking run.iterchildren() in order and
        emitting a real <br/> for every <w:br/> preserves the same forced
        wrap point the DOCX shows.
        """
        pieces = []
        for child in run:
            ctag = child.tag.split('}')[-1]
            if ctag == 't':
                pieces.append(child.text or "")
            elif ctag == 'br':
                pieces.append("<br/>")
            elif ctag == 'tab':
                pieces.append("\t")
        return "".join(pieces)

    def _runs_to_html(para_element) -> str:
        parts = []
        for run in para_element.findall(qn('w:r')):
            run_text = _run_inner_html(run)
            if not run_text:
                continue
            rpr = run.find(qn('w:rPr'))
            is_underline = (
                rpr is not None and rpr.find(qn('w:u')) is not None and
                rpr.find(qn('w:u')).get(qn('w:val'), 'single') != 'none'
            )
            if is_underline:
                run_text = f"<u>{run_text}</u>"
            if preserve_style and rpr is not None and rpr.find(qn('w:b')) is not None:
                b_val = rpr.find(qn('w:b')).get(qn('w:val'))
                if b_val is None or b_val not in ("false", "0"):
                    run_text = f"<strong>{run_text}</strong>"
            parts.append(run_text)
        return "".join(parts)

    def _runs_to_html_plain(para_element) -> str:
        """Exact original (preserve_style=False) cell/paragraph run logic —
        underline only, no bold, no alignment."""
        parts = []
        for run in para_element.findall('.//' + qn('w:r')):
            run_text = _run_inner_html(run)
            if not run_text:
                continue
            rpr = run.find(qn('w:rPr'))
            is_underline = (
                rpr is not None and rpr.find(qn('w:u')) is not None and
                rpr.find(qn('w:u')).get(qn('w:val'), 'single') != 'none'
            )
            parts.append(f"<u>{run_text}</u>" if is_underline else run_text)
        return "".join(parts)

    def _row_is_bold(row_elem) -> bool:
        """
        Heuristic for whether a table row is a genuine header row: true only
        if the row has actual text AND every run with text is bold.

        None of item001/item002's tables have a real Word tblHeader marker
        (checked: zero <w:tblHeader> elements in either file) — whatever
        generated these docs treats row 0 as a markdown-style header
        whenever the table happens to have a "|---|" separator line, even
        for tables that aren't conceptually headed at all (the filer-status
        checkbox table, the TOC's item-listing tables). Position isn't a
        reliable signal here; whether the row is actually bold-formatted in
        the source is — every table where row 0 SHOULD look like a header
        (state/EIN, securities) has real bold runs there, and every table
        where it shouldn't (filer checkboxes, TOC) doesn't.
        """
        saw_text = False
        for cell in row_elem.findall('.//' + qn('w:tc')):
            for run in cell.findall('.//' + qn('w:r')):
                run_text = "".join(t.text or "" for t in run.findall('.//' + qn('w:t')))
                if not run_text.strip():
                    continue
                saw_text = True
                rpr = run.find(qn('w:rPr'))
                b_elem = rpr.find(qn('w:b')) if rpr is not None else None
                is_bold = b_elem is not None and b_elem.get(qn('w:val')) not in ("false", "0")
                if not is_bold:
                    return False
        return saw_text

    def _cell_has_border(cell_elem) -> bool:
        """
        True only if the cell's own tcBorders actually specifies a visible
        border side (val not 'nil'/'none'). Confirmed by inspecting the raw
        XML: item001's tables have real w:val="single" borders; item002's
        <w:tcBorders></w:tcBorders> is empty (no child elements at all) —
        i.e. TOC tables genuinely have no borders in the source, so nothing
        should be added for them.
        """
        tcpr = cell_elem.find(qn('w:tcPr'))
        if tcpr is None:
            return False
        borders = tcpr.find(qn('w:tcBorders'))
        if borders is None:
            return False
        for side in ('top', 'start', 'left', 'bottom', 'end', 'right'):
            side_elem = borders.find(qn(f'w:{side}'))
            if side_elem is not None and side_elem.get(qn('w:val')) not in (None, "nil", "none"):
                return True
        return False

    def _cell_shading_color(cell_elem) -> str:
        """
        This cell's w:shd fill color (hex, no '#'), or '' if none set.
        Only the four financial-statement table writers
        (_data_row/_total_row/_section_header_row and the Stockholders'
        Equity writer) ever call _add_cell_shading() at docx-generation
        time — notes tables, TOC tables, and the cover page are built
        through entirely different code paths and never get w:shd at all,
        so this returns '' for every one of their cells regardless of
        table shape.
        """
        tcpr = cell_elem.find(qn('w:tcPr'))
        if tcpr is None:
            return ""
        shd = tcpr.find(qn('w:shd'))
        if shd is None:
            return ""
        fill = shd.get(qn('w:fill'), "")
        return fill if fill and fill.lower() != "auto" else ""

    def _cell_colspan(cell_elem) -> int:
        """
        This cell's horizontal span (w:gridSpan), or 1 if unset.

        _section_header_row() merges its label cell across every column
        via python-docx's cell.merge() — which, in the underlying OOXML,
        collapses the merged columns into a SINGLE <w:tc> carrying
        w:gridSpan, not multiple <w:tc> elements. Nothing downstream of
        that read gridSpan back out, so a 3-column section-header row
        ("EQUITY", "ASSETS", etc.) was emitted as exactly ONE <td>/<th>
        with no colspan — visually correct while the row had no
        background, but once _add_row_shading() gave these rows a real
        background color, the gap where columns 2/3 have no cell at all
        stopped taking that color, since a browser only paints the
        background behind actual cell boxes, not empty grid space beyond
        the last real cell in a row.
        """
        tcpr = cell_elem.find(qn('w:tcPr'))
        if tcpr is None:
            return 1
        span = tcpr.find(qn('w:gridSpan'))
        if span is None:
            return 1
        try:
            return int(span.get(qn('w:val'), "1"))
        except (TypeError, ValueError):
            return 1

    # Matches Amazon's own cover-page convention (10pt Times New Roman body
    # text) — added as a REAL inline style, not a <style> rule, for the same
    # EFM 5.02.05 reason as the heading-centering/border fixes above.
    _BODY_FONT = "font-family: 'Times New Roman', serif; font-size: 10pt;"
    _HEADING_FONT = {1: "font-family: 'Times New Roman', serif; font-size: 14pt; font-weight: bold;",
                     2: "font-family: 'Times New Roman', serif; font-size: 12pt; font-weight: bold;",
                     3: "font-family: 'Times New Roman', serif; font-size: 11pt; font-weight: bold;"}

    html_parts = []
    for element in doc.element.body:
        tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag

        if tag == 'p':
            style_elem = element.find(qn('w:pPr'))
            style_name = ""
            if style_elem is not None:
                style_ref = style_elem.find(qn('w:pStyle'))
                if style_ref is not None:
                    style_name = style_ref.get(qn('w:val'), "")
            align = _JC_MAP.get(_para_jc(element), "")
            text_content = (
                _runs_to_html(element) if preserve_style else
                "".join(t.text or "" for t in element.iter() if t.tag.split('}')[-1] == 't')
            )
            # Headings always get centered + a real font declaration in
            # preserve_style mode via inline style — not a <style> block.
            # EDGAR strips <style> blocks from the filed .htm (EFM 5.02.05:
            # no inline stylesheets), so a global CSS rule never survives
            # into the actual filed document. Per-element inline style="..."
            # is the only thing that does.
            if "Heading1" in style_name or "heading1" in style_name.lower():
                h_style = f' style="text-align: center; {_HEADING_FONT[1]}"' if preserve_style else ""
                html_parts.append(f"<h1{h_style}>{text_content}</h1>")
            elif "Heading2" in style_name or "heading2" in style_name.lower():
                h_style = f' style="text-align: center; {_HEADING_FONT[2]}"' if preserve_style else ""
                html_parts.append(f"<h2{h_style}>{text_content}</h2>")
            elif "Heading3" in style_name or "heading3" in style_name.lower():
                h_style = f' style="text-align: center; {_HEADING_FONT[3]}"' if preserve_style else ""
                html_parts.append(f"<h3{h_style}>{text_content}</h3>")
            elif text_content.strip():
                if preserve_style:
                    p_style = _BODY_FONT + (f" text-align: {align};" if align else "")
                    html_parts.append(f'<p style="{p_style}">{text_content}</p>')
                else:
                    html_parts.append(f"<p{_para_align_style(element)}>{text_content}</p>")

        elif tag == 'tbl':
            table_style = ' style="border-collapse: collapse; width: 100%"' if preserve_style else ""
            html_parts.append(f"<table{table_style}>")

            rows_xml = element.findall('.//' + qn('w:tr'))
            # An "Item N." listing table (TOC-style: Item 1., Item 1A., ...)
            # needs a narrow first column — otherwise table-layout:auto with
            # width:100% stretches it to share space evenly with the long
            # description column next to it. Detected by content pattern,
            # not column count, so the securities table (also 3 columns,
            # but its first column IS a long value — "Common Stock, without
            # par value" — and should NOT be narrowed) is left alone.
            is_item_listing_table = False
            if preserve_style and rows_xml:
                import re as _re
                first_col_texts = []
                for row in rows_xml:
                    cells = row.findall('.//' + qn('w:tc'))
                    if cells:
                        txt = "".join(
                            t.text or "" for t in cells[0].iter()
                            if t.tag.split('}')[-1] == 't'
                        ).strip()
                        first_col_texts.append(txt)
                is_item_listing_table = bool(first_col_texts) and all(
                    _re.match(r'^Item\s+\d+[A-Za-z]?\.?$', t) or t == ""
                    for t in first_col_texts
                ) and any(first_col_texts)

            for row_idx, row in enumerate(rows_xml):
                row_cells_xml = row.findall('.//' + qn('w:tc'))
                row_bg = ""
                if preserve_style and row_cells_xml:
                    shades = [_cell_shading_color(c) for c in row_cells_xml]
                    # Only emit a row-level background if EVERY cell in the
                    # row carries the SAME shading — true for every row
                    # _add_row_shading() touched (it always shades the
                    # whole row uniformly), and never true for notes/TOC/
                    # cover tables, whose cells never have w:shd set at all
                    # (shades[0] == "" there, so this is skipped).
                    if shades[0] and all(s == shades[0] for s in shades):
                        row_bg = shades[0]
                tr_style = f' style="background-color: #{row_bg};"' if row_bg else ""
                html_parts.append(f"<tr{tr_style}>")
                is_header = _row_is_bold(row) if preserve_style else _is_header_row(row, element)
                for col_idx, cell in enumerate(row.findall('.//' + qn('w:tc'))):
                    para_texts = []
                    cell_align_style = ""
                    for para in cell.findall('.//' + qn('w:p')):
                        if not cell_align_style:
                            cell_align_style = _para_align_style(para)
                        para_html = _runs_to_html(para) if preserve_style else _runs_to_html_plain(para)
                        if para_html:
                            para_texts.append(para_html)
                    cell_text = "<br/>".join(para_texts)
                    plain_text = cell_text.replace("<u>", "").replace("</u>", "") \
                                            .replace("<strong>", "").replace("</strong>", "")
                    tag_name = "th" if is_header else "td"
                    css_class = ' class="number"' if _looks_like_number(plain_text) else ''
                    if preserve_style:
                        decls = [_BODY_FONT, "padding: 4px 8px;"]
                        if _cell_has_border(cell):
                            decls.append("border: 1px solid #000;")
                        align_decl = _JC_MAP.get(_cell_jc(cell), "")
                        if align_decl:
                            decls.append(f"text-align: {align_decl};")
                        if is_item_listing_table and col_idx == 0:
                            decls.append("width: 8%; white-space: nowrap;")
                        style_attr = f' style="{" ".join(decls)}"'
                    else:
                        style_attr = cell_align_style
                    colspan = _cell_colspan(cell)
                    colspan_attr = f' colspan="{colspan}"' if colspan > 1 else ""
                    html_parts.append(
                        f"<{tag_name}{css_class}{colspan_attr}{style_attr}>{cell_text}</{tag_name}>"
                    )
                html_parts.append("</tr>")
            html_parts.append("</table>")

    return html_parts


def convert_to_edgar_html(docx_path: str, output_filename: str) -> str:
    """Convert a DOCX file to SEC EDGAR-compliant HTML.
    Writes to REPORTS_DIR (container path), bind-mounted to WIN_REPORTS_DIR.

    LEGACY PLAIN PATH — kept for reference / non-EDGAR previews of a single
    docx file. Still emits the exact original bare <p>/<th>/<td> markup
    (preserve_style=False), with only blanket tag-type-level CSS in the
    shared <style> block (table/cell borders, h2/h3 centering) — no
    tag/attribute-shape changes.

    The final 10-K's tagged EDGAR HTML no longer goes through this
    function — create_10k_edgar_html() now builds BOTH the intro and body
    fragments with preserve_style=True (real inline styles, since EDGAR
    strips <style> blocks per EFM 5.02.05). This function remains
    available if a plain, unstyled conversion is ever needed again.
    """
    _ensure_dir(settings.REPORTS_DIR)
    output_path = os.path.join(settings.REPORTS_DIR, output_filename)

    doc = Document(docx_path)
    html_parts = [
        "<?xml version='1.0' encoding='UTF-8'?>",
        "<!DOCTYPE html PUBLIC \"-//W3C//DTD XHTML 1.0 Strict//EN\" "
        "\"http://www.w3.org/TR/xhtml1/DTD/xhtml1-strict.dtd\">",
        "<html xmlns='http://www.w3.org/1999/xhtml'>",
        "<head>",
        "<meta http-equiv='Content-Type' content='text/html; charset=UTF-8'/>",
        f"<title>{output_filename.replace('.htm', '')}</title>",
        "<style type='text/css'>", _EDGAR_CSS, "</style>",
        "</head>",
        "<body>",
    ]
    html_parts.extend(_docx_body_to_html_parts(doc, preserve_style=False))
    html_parts.extend(["</body>", "</html>"])
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(html_parts))
    return output_path


def _is_header_row(row_elem, tbl_elem) -> bool:
    from docx.oxml.ns import qn
    all_rows = tbl_elem.findall('.//' + qn('w:tr'))
    return len(all_rows) > 0 and all_rows[0] == row_elem


def _looks_like_number(text: str) -> bool:
    cleaned = text.replace("$", "").replace(",", "").replace("(", "").replace(")", "").replace("-", "").strip()
    try:
        float(cleaned)
        return True
    except ValueError:
        return False


# ─── Styled cover/TOC fragment + full-document HTML merge ────────────────────

def convert_intro_to_html(docx_paths: list) -> str:
    """
    Convert item001-10-k-cover.docx and item002-10-k-TOC.docx into a single
    styled HTML fragment (centering + bold preserved), for splicing into the
    front of the final EDGAR HTML by create_10k_edgar_html().

    This is the ONLY place that produces styled (non-bare) cover markup.
    tag_cover_page() in xbrl_tagger.py is written to tolerate exactly the
    two things this can add to a tag: an optional inline
    style="text-align: ..." attribute on the opening tag, and an optional
    <strong>...</strong> wrapper around bold text. If you change what this
    function emits, re-check tag_cover_page()'s regexes against the new
    output before trusting Arelle validation again.
    """
    fragment_parts = []
    for path in docx_paths:
        doc = Document(path)
        fragment_parts.extend(_docx_body_to_html_parts(doc, preserve_style=True))
    return "\n".join(fragment_parts)


def _merge_10k_body_only_docx() -> str:
    """
    Merge Part I-IV + Signatures into a temp docx — same as
    create_10k_final_report(), but WITHOUT the intro directory, since the
    intro (cover + TOC) is converted separately by convert_intro_to_html()
    and spliced in as styled HTML instead.
    """
    _ensure_dir(settings.REPORTS_DIR)
    tmp_path = os.path.join(settings.REPORTS_DIR, "_sec-10k-body-only.docx")

    master = Document()
    composer = Composer(master, preserve_styles=True)
    _merge_dir_into_composer(composer, settings.DATA_10K_PART1, heading="PART I.")
    _merge_dir_into_composer(composer, settings.DATA_10K_PART2, heading="PART II.")
    _merge_dir_into_composer(composer, settings.DATA_10K_PART3, heading="PART III.")
    _merge_dir_into_composer(composer, settings.DATA_10K_PART4, heading="PART IV.")
    master.add_page_break()
    _merge_dir_into_composer(composer, settings.DATA_10K_SIGS)

    master.save(tmp_path)
    return tmp_path


def _reset_h3_align_left(body_html: str) -> str:
    """
    Force every <h3> heading in the BODY fragment back to left-aligned.

    _docx_body_to_html_parts(preserve_style=True) hardcodes
    `text-align: center` on every heading level (h1/h2/h3) in its inline
    style, regardless of the source docx paragraph's own alignment — see
    that function's heading-handling comment. That's correct for
    item001/002's cover-page headings (Heading 1/Heading 2 — genuinely
    centered in the source docx), but every Heading 3 paragraph in the
    BODY (Part headings, Item headings, "Notes to Consolidated Financial
    Statements", each note's/MD&A section's own <h3> heading, etc.) is
    LEFT-aligned in Ron's source docx files.

    Rather than threading the real per-paragraph alignment through the
    shared h1/h2/h3 branch (which also serves the genuinely-centered
    intro headings, and would need to distinguish the two call sites),
    this runs as a single, narrowly-scoped BeautifulSoup post-processing
    pass over just the body fragment — never the intro — overriding every
    <h3>'s inline text-align back to left. h1/h2 (used only by the
    cover/TOC intro fragment, which this function is never called on) are
    untouched by construction.

    Note: parsing+reserializing through BeautifulSoup also normalizes any
    bare "&" in the body text to "&amp;" (e.g. "R&D Expense" ->
    "R&amp;D Expense"). This is a side effect of the round-trip, not a
    bug — the filing's own DOCTYPE declares strict XHTML 1.0, where a
    bare "&" is technically invalid, and BeautifulSoup's .get_text() (used
    everywhere financial-statement/note tagging reads cell or paragraph
    text — see _tag_table, tag_notes_section) transparently unescapes it
    back, so no downstream tagging logic sees any difference.
    """
    soup = BeautifulSoup(body_html, "html.parser")
    for h3 in soup.find_all("h3"):
        style = h3.get("style", "")
        if "text-align" in style:
            style = re.sub(r"text-align\s*:\s*[^;]+;?", "text-align: left;", style)
        else:
            style = (style.rstrip() + " text-align: left;").strip()
        h3["style"] = style
    return str(soup)


_TOC_PART_RE = re.compile(r'^\s*PART\s+([IVXLCDM]+)\b', re.IGNORECASE)
_TOC_ITEM_RE = re.compile(r'^\s*Item\s+(\d+[A-Za-z]?)\.', re.IGNORECASE)


def _add_toc_navigation_links(combined_html: str) -> str:
    """
    Give every "Item N." / part heading in the BODY an id="..." anchor, and
    turn the matching row in the Table of Contents (in the intro fragment)
    into a link to that anchor — so a reader can click any TOC entry and
    jump straight to the section, the way real EDGAR filings do. No page
    numbers are involved; this is pure in-document navigation.

    Must run on the INTRO + BODY fragments concatenated together (both are
    needed — anchors are assigned in the body, links point at them from
    the intro) but never on the full assembled document (with its
    DOCTYPE/meta/<style> block): BeautifulSoup's html.parser round-trip
    normalizes things like self-closing tags, which is fine for a bare
    fragment (see _reset_h3_align_left's docstring for the same
    reasoning) but must not be risked against the parts of the document
    EDGAR/Arelle parsing is strict about.

    10-Q filings reuse "Item 1.", "Item 2.", etc. under BOTH Part I and
    Part II, so anchors are qualified by the most recently seen "PART n"
    heading (e.g. part-i-item-1 vs part-ii-item-1) rather than by item
    number alone.
    """
    soup = BeautifulSoup(combined_html, "html.parser")

    # ---- Pass 1: body headings — assign ids ----
    current_part = None
    slug_for = {}  # (part_token, item_token) -> slug, in first-seen order

    for el in soup.find_all(['h3', 'p']):
        if el.find_parent('table') is not None:
            continue  # TOC rows are handled in Pass 2
        text = el.get_text().strip()
        if not text:
            continue

        if el.name == 'h3':
            part_match = _TOC_PART_RE.match(text)
            if part_match:
                current_part = part_match.group(1).upper()
                continue
            if text.upper().startswith('SIGNATURES'):
                el['id'] = 'signatures'
                continue

        item_match = _TOC_ITEM_RE.match(text)
        if item_match:
            item_token = item_match.group(1).upper()
            key = (current_part, item_token)
            if key not in slug_for:
                slug = f"part-{(current_part or 'x').lower()}-item-{item_token.lower()}"
                el['id'] = slug
                slug_for[key] = slug

    # ---- Pass 2: TOC tables — wrap matching cells in <a href="#slug"> ----
    current_part = None
    for el in soup.find_all(['p', 'table']):
        if el.name == 'p':
            text = el.get_text().strip()
            part_match = _TOC_PART_RE.match(text)
            if part_match:
                current_part = part_match.group(1).upper()
            continue

        rows = el.find_all('tr')
        if not rows:
            continue
        first_cells = [r.find(['td', 'th']) for r in rows]
        if not all(fc is not None for fc in first_cells):
            continue
        first_texts = [fc.get_text().strip() for fc in first_cells]
        looks_like_toc = bool(first_texts) and any(first_texts) and all(
            _TOC_ITEM_RE.match(t) or t == "" for t in first_texts
        )
        if not looks_like_toc:
            continue

        for row in rows:
            cells = row.find_all(['td', 'th'])
            if len(cells) < 2:
                continue
            item_cell, desc_cell = cells[0], cells[1]
            item_text = item_cell.get_text().strip()
            desc_text = desc_cell.get_text().strip()

            if item_text == "" and desc_text.upper().startswith("SIGNATURES"):
                slug = "signatures"
            else:
                item_match = _TOC_ITEM_RE.match(item_text)
                if not item_match:
                    continue
                slug = slug_for.get((current_part, item_match.group(1).upper()))
                if not slug:
                    continue

            for cell in (item_cell, desc_cell):
                if not cell.get_text().strip():
                    continue
                a_tag = soup.new_tag(
                    'a', href=f"#{slug}",
                    style="color: #0000ff; text-decoration: underline;"
                )
                for child in list(cell.contents):
                    a_tag.append(child.extract())
                cell.append(a_tag)

    return str(soup)


def create_10k_edgar_html(output_filename: str = "sec-10k.htm") -> str:
    """
    Build the final EDGAR HTML by merging TWO separately-generated pieces,
    BOTH now going through the SAME styled walker
    (_docx_body_to_html_parts(preserve_style=True)):
      1. item001 (cover) + item002 (TOC), via convert_intro_to_html().
      2. Parts I-IV + Signatures (financial statements, Notes, MD&A, etc.).

    Styled markup means every paragraph carries an inline
    style="font-family: ...; font-size: ...[; text-align: ...]" attribute,
    headings get a real font/centering declaration, and bold runs are
    wrapped in <strong>...</strong> — all as REAL inline attributes, never
    a <style> block, since EDGAR strips <style> blocks entirely (EFM
    5.02.05) and inline style="..." is the only presentational information
    that survives into the filed .htm. See _docx_body_to_html_parts()'s
    docstring for the full preserve_style=True behavior.

    This is a change from the prior behavior, where only the intro
    fragment was styled and the body (including Notes to Financial
    Statements) came out as bare, unstyled <p>/<table> markup that
    rendered with zero formatting once EDGAR stripped the <style> block.

    xbrl_tagger.py's financial-statement table tagging (_tag_table /
    _tag_stockholders_equity_table) is unaffected by this switch — it
    walks the parsed BeautifulSoup tree and reads cell text via
    .get_text(), which is blind to inline style attributes and <strong>
    wrappers either way. The raw-string regex passes that run BEFORE
    BeautifulSoup parses the document (tag_cover_page,
    tag_auditor_report_block, tag_notes_section, and notes_config.py's
    heading/end patterns) DO care about the exact markup shape, and have
    been made tolerant of an optional inline style="..." attribute and an
    optional <strong> wrapper — see each of those functions' own
    docstrings/comments for the specific tolerance added.

    Call fill_cover_page(period_end) before this so item001 reflects the
    current business_info.toml. xbrl_tagger.tag_filing() runs on the
    output of THIS function (not convert_to_edgar_html(), which remains
    the plain, unstyled legacy conversion path) for the final 10-K.
    """
    _ensure_dir(settings.REPORTS_DIR)
    output_path = os.path.join(settings.REPORTS_DIR, output_filename)

    intro_paths = sorted(Path(settings.DATA_10K_INTRO).glob("*.docx"))
    intro_html = convert_intro_to_html([str(p) for p in intro_paths])

    body_docx_path = _merge_10k_body_only_docx()
    body_doc = Document(body_docx_path)
    body_parts = _docx_body_to_html_parts(body_doc, preserve_style=True)
    # Every <h3> comes out of _docx_body_to_html_parts hardcoded to
    # text-align: center (see that function's heading branch) — correct
    # for the intro's Heading 1/2, but Ron's actual Heading 3 paragraphs
    # throughout this body are left-aligned in the source docx. Reset
    # just this fragment's h3 alignment back to left — see
    # _reset_h3_align_left's docstring.
    body_html = _reset_h3_align_left("\n".join(body_parts))

    # TOC <-> section jump links: needs the intro (TOC) and body (section
    # headings) together in one pass, but must stay OUTSIDE the
    # DOCTYPE/meta/<style> wrapper below — see
    # _add_toc_navigation_links()'s docstring.
    linked_intro_and_body = _add_toc_navigation_links(intro_html + "\n" + body_html)

    html_parts = [
        "<?xml version='1.0' encoding='UTF-8'?>",
        "<!DOCTYPE html PUBLIC \"-//W3C//DTD XHTML 1.0 Strict//EN\" "
        "\"http://www.w3.org/TR/xhtml1/DTD/xhtml1-strict.dtd\">",
        "<html xmlns='http://www.w3.org/1999/xhtml'>",
        "<head>",
        "<meta http-equiv='Content-Type' content='text/html; charset=UTF-8'/>",
        f"<title>{output_filename.replace('.htm', '')}</title>",
        "<style type='text/css'>", _EDGAR_CSS, "</style>",
        "</head>",
        "<body>",
        linked_intro_and_body,
    ]
    html_parts.extend(["</body>", "</html>"])

    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(html_parts))

    try:
        os.remove(body_docx_path)
    except OSError:
        pass

    return output_path
